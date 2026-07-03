from __future__ import annotations

import argparse
import math
import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from PIL import Image, ImageDraw, ImageFont


TABLE_STYLE = "Table Grid"
HEADING_STYLES = tuple(f"Heading {level}" for level in range(1, 5))
EXCLUDED_MARKDOWN_FILENAMES = {"agents.md", "readme.md"}
BASE_FONT = "Calibri"
BASE_FONT_SIZE = Pt(11)
LIST_LEFT_INDENT = Inches(0.25)
LIST_TEXT_TAB = Inches(0.25)
LIST_LEFT_DXA = "360"
LIST_HANGING_DXA = "360"


def read_text(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def strip_front_matter(text: str) -> str:
    lines = text.splitlines()
    if lines and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            if lines[index].strip() == "---":
                return "\n".join(lines[index + 1 :]).lstrip()
    return text


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"Imagen: \1 (\2)", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    return text.strip()


def has_style(document: Document, name: str) -> bool:
    try:
        document.styles[name]
    except KeyError:
        return False
    return True


def set_paragraph_format(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_after = Pt(0)


def set_list_paragraph_format(paragraph) -> None:
    set_paragraph_format(paragraph)
    paragraph.paragraph_format.left_indent = LIST_LEFT_INDENT
    paragraph.paragraph_format.first_line_indent = -LIST_LEFT_INDENT
    paragraph.paragraph_format.tab_stops.add_tab_stop(LIST_TEXT_TAB)


def set_run_font(run, font_name: str = BASE_FONT, font_size=BASE_FONT_SIZE) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = font_size


def apply_base_styles(document: Document) -> None:
    default_document = Document()
    for name in ("Normal", "List Paragraph", "List Bullet", "List Number", *HEADING_STYLES):
        if not has_style(document, name) and has_style(default_document, name):
            document.styles.element.append(deepcopy(default_document.styles[name]._element))

        if not has_style(document, name):
            continue

        style = document.styles[name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        style.font.size = BASE_FONT_SIZE
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_after = Pt(0)


def next_numbering_id(numbering, element_name: str, attr_name: str, default: int) -> int:
    values: list[int] = []
    for element in numbering.iter(qn(f"w:{element_name}")):
        value = element.get(qn(f"w:{attr_name}"))
        if value is not None and value.isdigit():
            values.append(int(value))
    return max(values + [default - 1]) + 1


def ensure_bullet_numbering(document: Document) -> str:
    numbering = document.part.numbering_part.element
    abstract_num_id = str(next_numbering_id(numbering, "abstractNum", "abstractNumId", 0))
    num_id = str(next_numbering_id(numbering, "num", "numId", 1))

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_num_id)

    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level_type)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet")
    level.append(num_format)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022")
    level.append(level_text)

    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.append(level_justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), LIST_LEFT_DXA)
    tabs.append(tab)
    paragraph_properties.append(tabs)

    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), LIST_LEFT_DXA)
    indent.set(qn("w:hanging"), LIST_HANGING_DXA)
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    abstract_num.append(level)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_num_ref = OxmlElement("w:abstractNumId")
    abstract_num_ref.set(qn("w:val"), abstract_num_id)
    num.append(abstract_num_ref)
    numbering.append(num)

    return num_id


def apply_numbering(paragraph, num_id: str, level: str = "0") -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    existing_numbering = paragraph_properties.find(qn("w:numPr"))
    if existing_numbering is not None:
        paragraph_properties.remove(existing_numbering)

    numbering_properties = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), level)
    numbering_properties.append(ilvl)

    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), num_id)
    numbering_properties.append(num_id_element)
    paragraph_properties.append(numbering_properties)


def ensure_heading_styles(document: Document) -> None:
    default_document = Document()
    for name in HEADING_STYLES:
        if not has_style(document, name):
            document.styles.element.append(
                deepcopy(default_document.styles[name]._element)
            )


def is_separator_row(line: str) -> bool:
    cells = split_table_row(line)
    if not cells:
        return False
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return "|" in lines[index] and is_separator_row(lines[index + 1])


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(cell.strip()) for cell in line.split("|")]


def collect_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = [split_table_row(lines[index])]
    index += 2
    while index < len(lines):
        line = lines[index]
        if not line.strip() or "|" not in line:
            break
        rows.append(split_table_row(line))
        index += 1
    return rows, index


def is_field_value_table(rows: list[list[str]]) -> bool:
    if not rows or len(rows[0]) != 2:
        return False
    first_cell = rows[0][0].strip().casefold()
    second_cell = rows[0][1].strip().casefold()
    return first_cell == "campo" and second_cell == "valor"


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    field_value_table = is_field_value_table(rows)
    if field_value_table:
        rows = rows[1:]
        if not rows:
            return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = TABLE_STYLE
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        for column_index in range(column_count):
            value = row_values[column_index] if column_index < len(row_values) else ""
            cell = row.cells[column_index]
            cell.text = value
            if row_index == 0 and not field_value_table:
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph)
                    for run in paragraph.runs:
                        run.bold = True
                        set_run_font(run)
            elif field_value_table and column_index == 0:
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph)
                    for run in paragraph.runs:
                        run.bold = True
                        set_run_font(run)


def add_code_block(document: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run("\n".join(code_lines))
    set_run_font(run, "Consolas", Pt(9))


def load_diagram_font(size: int, bold: bool = False):
    font_names = ["calibrib.ttf", "calibri.ttf"] if bold else ["calibri.ttf"]
    for font_name in font_names:
        font_path = Path("C:/Windows/Fonts") / font_name
        if font_path.exists():
            return ImageFont.truetype(str(font_path), size)
    return ImageFont.load_default()


def clean_mermaid_label(value: str) -> str:
    value = value.strip()
    if len(value) >= 2 and value[0] == value[-1] == '"':
        value = value[1:-1]
    return value.replace("<br/>", "\n").replace("<br>", "\n")


def parse_mermaid_endpoint(value: str) -> tuple[str, str, str]:
    value = value.strip()
    match = re.match(
        r'^([A-Za-z][\w-]*)(?:\[(?:"([^"]*)"|([^\]]*))\]|\{(?:"([^"]*)"|([^}]*))\})?$',
        value,
    )
    if not match:
        return value, value, "process"

    node_id = match.group(1)
    square_label = match.group(2) or match.group(3)
    diamond_label = match.group(4) or match.group(5)
    if diamond_label is not None:
        return node_id, clean_mermaid_label(diamond_label), "decision"
    if square_label is not None:
        return node_id, clean_mermaid_label(square_label), "process"
    return node_id, node_id, "process"


def parse_mermaid_flowchart(code_lines: list[str]):
    nodes: dict[str, dict[str, str]] = {}
    edges: list[tuple[str, str, str]] = []

    edge_pattern = re.compile(
        r'(.+?)\s+(?:--\s*"([^"]+)"\s*-->|-->|-\.\s*"([^"]+)"\s*\.->|-\.->)\s+(.+)$'
    )

    for raw_line in code_lines:
        line = raw_line.strip()
        if not line or line.startswith("flowchart ") or line.startswith("graph "):
            continue

        edge = edge_pattern.match(line)
        if not edge:
            continue

        source_raw, label_a, label_b, target_raw = edge.groups()
        source_id, source_label, source_shape = parse_mermaid_endpoint(source_raw)
        target_id, target_label, target_shape = parse_mermaid_endpoint(target_raw)

        for node_id, node_label, node_shape in (
            (source_id, source_label, source_shape),
            (target_id, target_label, target_shape),
        ):
            if node_id not in nodes or node_label != node_id:
                nodes[node_id] = {"label": node_label, "shape": node_shape}
        edges.append((source_id, target_id, clean_mermaid_label(label_a or label_b or "")))

    return nodes, edges


def mermaid_node_levels(nodes: dict[str, dict[str, str]], edges: list[tuple[str, str, str]]):
    incoming = {node_id: 0 for node_id in nodes}
    outgoing: dict[str, list[str]] = {node_id: [] for node_id in nodes}
    for source, target, _label in edges:
        incoming[target] = incoming.get(target, 0) + 1
        outgoing.setdefault(source, []).append(target)

    roots = [node_id for node_id, count in incoming.items() if count == 0] or list(nodes)[:1]
    levels = {node_id: math.inf for node_id in nodes}
    queue: list[str] = []
    for root in roots:
        levels[root] = 0
        queue.append(root)

    while queue:
        current = queue.pop(0)
        for target in outgoing.get(current, []):
            next_level = levels[current] + 1
            if next_level < levels[target]:
                levels[target] = next_level
                queue.append(target)

    max_level = max((level for level in levels.values() if level != math.inf), default=0)
    for node_id, level in list(levels.items()):
        if level == math.inf:
            max_level += 1
            levels[node_id] = max_level
    return levels


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.splitlines() or [""]:
        words = paragraph.split()
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines or [text]


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font,
    fill: str,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(80, x2 - x1 - 40))
    line_height = font.size + 7 if hasattr(font, "size") else 18
    total_height = len(lines) * line_height
    y = y1 + ((y2 - y1 - total_height) / 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + ((x2 - x1 - (bbox[2] - bbox[0])) / 2)
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str,
) -> None:
    draw.line([start, end], fill=color, width=3)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 13
    left = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def render_mermaid_flowchart(code_lines: list[str]) -> BytesIO | None:
    nodes, edges = parse_mermaid_flowchart(code_lines)
    if not nodes or not edges:
        return None

    levels = mermaid_node_levels(nodes, edges)
    rows: dict[int, list[str]] = {}
    for node_id in nodes:
        rows.setdefault(int(levels[node_id]), []).append(node_id)

    node_width = 420
    node_height = 118
    x_gap = 80
    y_gap = 105
    margin = 70
    max_row_count = max(len(row) for row in rows.values())
    width = max(1200, margin * 2 + max_row_count * node_width + (max_row_count - 1) * x_gap)
    height = margin * 2 + len(rows) * node_height + (len(rows) - 1) * y_gap

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font = load_diagram_font(27)
    small_font = load_diagram_font(22)

    positions: dict[str, tuple[int, int, int, int]] = {}
    for row_index, level in enumerate(sorted(rows)):
        row = rows[level]
        row_width = len(row) * node_width + (len(row) - 1) * x_gap
        x = (width - row_width) // 2
        y = margin + row_index * (node_height + y_gap)
        for node_id in row:
            positions[node_id] = (x, y, x + node_width, y + node_height)
            node = nodes[node_id]
            fill = "#F8FAFC" if node["shape"] == "process" else "#FFF7ED"
            outline = "#64748B" if node["shape"] == "process" else "#B45309"
            if node["shape"] == "decision":
                cx = x + node_width // 2
                cy = y + node_height // 2
                points = [(cx, y), (x + node_width, cy), (cx, y + node_height), (x, cy)]
                draw.polygon(points, fill=fill, outline=outline)
                draw.line(points + [points[0]], fill=outline, width=3)
            else:
                draw.rounded_rectangle(
                    [x, y, x + node_width, y + node_height],
                    radius=18,
                    fill=fill,
                    outline=outline,
                    width=3,
                )
            draw_centered_text(draw, positions[node_id], node["label"], font, "#111827")
            x += node_width + x_gap

    for source, target, label in edges:
        if source not in positions or target not in positions:
            continue
        source_box = positions[source]
        target_box = positions[target]
        start = ((source_box[0] + source_box[2]) // 2, source_box[3])
        end = ((target_box[0] + target_box[2]) // 2, target_box[1])
        draw_arrow(draw, start, end, "#334155")
        if label:
            mid_x = (start[0] + end[0]) // 2
            mid_y = (start[1] + end[1]) // 2
            bbox = draw.textbbox((0, 0), label, font=small_font)
            pad = 7
            draw.rounded_rectangle(
                [
                    mid_x - (bbox[2] - bbox[0]) // 2 - pad,
                    mid_y - (bbox[3] - bbox[1]) // 2 - pad,
                    mid_x + (bbox[2] - bbox[0]) // 2 + pad,
                    mid_y + (bbox[3] - bbox[1]) // 2 + pad,
                ],
                radius=8,
                fill="white",
                outline="#CBD5E1",
            )
            draw.text(
                (mid_x - (bbox[2] - bbox[0]) / 2, mid_y - (bbox[3] - bbox[1]) / 2),
                label,
                font=small_font,
                fill="#334155",
            )

    stream = BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    return stream


def add_mermaid_diagram(document: Document, code_lines: list[str]) -> bool:
    image_stream = render_mermaid_flowchart(code_lines)
    if image_stream is None:
        return False
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_stream, width=Inches(6.8))
    return True


def add_markdown_line(
    document: Document,
    line: str,
    center_if_title: bool = False,
    bullet_num_id: str | None = None,
) -> bool:
    stripped = line.strip()
    if not stripped:
        return False

    if re.fullmatch(r"[-*_]{3,}", stripped):
        return False

    heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 4)
        paragraph = document.add_heading(clean_inline(heading.group(2)), level=level)
        set_paragraph_format(paragraph)
        if center_if_title:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True

    unordered = re.match(r"^\s*[-*+]\s+(.+)$", line)
    if unordered:
        value = clean_inline(unordered.group(1))
        paragraph = document.add_paragraph(style="List Paragraph")
        set_list_paragraph_format(paragraph)
        if bullet_num_id is not None:
            apply_numbering(paragraph, bullet_num_id)
        run = paragraph.add_run(value)
        set_run_font(run)
        return True

    ordered = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
    if ordered:
        value = clean_inline(ordered.group(2))
        paragraph = document.add_paragraph(style="List Paragraph")
        set_list_paragraph_format(paragraph)
        run = paragraph.add_run(f"{ordered.group(1)}.\t{value}")
        set_run_font(run)
        return True

    quote = re.match(r"^>\s*(.+)$", stripped)
    if quote:
        paragraph = document.add_paragraph()
        set_paragraph_format(paragraph)
        run = paragraph.add_run(clean_inline(quote.group(1)))
        run.italic = True
        set_run_font(run)
        return True

    paragraph = document.add_paragraph(clean_inline(stripped))
    set_paragraph_format(paragraph)
    return True


def create_document(template: Path | None) -> Document:
    if template is None:
        document = Document()
        section = document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        return document

    document = Document(template)
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return document


def convert_markdown_to_docx(
    source: Path,
    output: Path,
    title_prefix: str,
    template: Path | None,
) -> None:
    text = strip_front_matter(read_text(source))
    lines = text.splitlines()

    document = create_document(template)
    ensure_heading_styles(document)
    apply_base_styles(document)
    bullet_num_id = ensure_bullet_numbering(document)

    document.core_properties.title = title_prefix

    index = 0
    in_code_block = False
    code_block_language = ""
    code_lines: list[str] = []
    first_content_pending = True

    while index < len(lines):
        line = lines[index]
        if line.strip().startswith("```"):
            if in_code_block:
                if code_block_language == "mermaid":
                    if not add_mermaid_diagram(document, code_lines):
                        add_code_block(document, code_lines)
                else:
                    add_code_block(document, code_lines)
                code_lines = []
                code_block_language = ""
                in_code_block = False
            else:
                in_code_block = True
                code_block_language = line.strip().strip("`").strip().lower()
            index += 1
            continue

        if in_code_block:
            code_lines.append(line)
            index += 1
            continue

        if is_table_start(lines, index):
            rows, index = collect_table(lines, index)
            add_table(document, rows)
            first_content_pending = False
            continue

        added_content = add_markdown_line(
            document,
            line,
            center_if_title=first_content_pending,
            bullet_num_id=bullet_num_id,
        )
        if added_content:
            first_content_pending = False
        index += 1

    if code_lines:
        if code_block_language == "mermaid":
            if not add_mermaid_diagram(document, code_lines):
                add_code_block(document, code_lines)
        else:
            add_code_block(document, code_lines)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def iter_markdown_files(source: Path, include_agents: bool) -> list[Path]:
    files = sorted(source.rglob("*.md"))
    return [
        path
        for path in files
        if path.name.lower() not in EXCLUDED_MARKDOWN_FILENAMES
    ]


def main() -> int:
    parser = argparse.ArgumentParser(description="Export Markdown files to Word documents.")
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--template", type=Path)
    parser.add_argument("--include-agents", action="store_true")
    args = parser.parse_args()

    source_root = args.source.resolve()
    output_root = args.output.resolve()
    template = args.template.resolve() if args.template else None

    if not source_root.exists():
        raise SystemExit(f"No existe la carpeta fuente: {source_root}")
    if template is not None and not template.exists():
        raise SystemExit(f"No existe la plantilla: {template}")

    exported = 0
    for markdown_path in iter_markdown_files(source_root, args.include_agents):
        relative = markdown_path.relative_to(source_root)
        output_path = output_root / source_root.name / relative.with_suffix(".docx")
        convert_markdown_to_docx(
            markdown_path,
            output_path,
            str(relative.with_suffix("")),
            template,
        )
        print(f"Exportado: {relative} -> {output_path.relative_to(output_root)}")
        exported += 1

    print("")
    print(f"Documentos exportados: {exported}")
    print(f"Carpeta de salida: {output_root}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
